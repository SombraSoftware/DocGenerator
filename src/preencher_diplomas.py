import os
from openpyxl import load_workbook
from docx import Document
from datetime import datetime


def mes_por_extenso(data):
    meses = [
        'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
        'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro'
    ]
    mes = int(data.strftime("%m"))
    return meses[mes - 1]


def substituir_tags_em_runs(paragraph, tags):
    """Substitui TAGs nos 'runs' de um parágrafo para preservar a formatação."""
    for run in paragraph.runs:
        for tag, valor in tags.items():
            if tag in run.text:
                run.text = run.text.replace(tag, valor)


def substituir_todas_as_tags(doc, tags):
    """Substitui todas as ocorrências das TAGs no documento sem alterar a formatação."""
    # Substituir TAGs nos parágrafos
    for paragraph in doc.paragraphs:
        substituir_tags_em_runs(paragraph, tags)

    # Substituir TAGs dentro das tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_tags_em_runs(paragraph, tags)


def preencher_certificados(planilha_excel, modelo_certificado, saida_certificados):
    # Carregar a planilha
    try:
        print("Carregando planilha...")
        workbook = load_workbook(planilha_excel)
        planilha = workbook.active
    except Exception as e:
        print(f"Erro ao carregar a planilha: {e}")
        return

    # Verificar se o diretório de saída existe, se não, criar
    if not os.path.exists(saida_certificados):
        os.makedirs(saida_certificados)

    # Para armazenar possíveis erros de preenchimento
    erros_preenchimento = []

    # Iterar sobre as linhas da planilha
    for row in range(2, planilha.max_row + 1):
        # Obter os dados da linha atual com a nova sequência de colunas
        nome_aluno = planilha.cell(row=row, column=1).value
        nascimento = planilha.cell(row=row, column=2).value
        cpf = planilha.cell(row=row, column=3).value
        cidade = planilha.cell(row=row, column=4).value
        estado = planilha.cell(row=row, column=5).value
        reg = planilha.cell(row=row, column=6).value
        sistec = planilha.cell(row=row, column=7).value

        # Verificar se o nome do aluno está vazio
        if nome_aluno is None:
            continue

        # Converter a data de nascimento para string com mês em português
        if isinstance(nascimento, datetime):
            nascimento_str = nascimento.strftime("%d de ") + mes_por_extenso(nascimento) + nascimento.strftime(" de %Y")
        else:
            nascimento_str = str(nascimento)

        # Criar o documento do Word baseado no modelo
        try:
            print(f"Preenchendo certificado para {nome_aluno}...")
            doc = Document(modelo_certificado)
        except Exception as e:
            print(f"Erro ao carregar o modelo de certificado: {e}")
            continue

        # Tags a serem substituídas e seus valores
        tags = {
            "{NOME_ALUNO}": nome_aluno,
            "{NASCIMENTO}": nascimento_str,
            "{CPF}": str(cpf) if cpf is not None else "",
            "{REG}": str(reg) if reg is not None else "",
            "{SISTEC}": str(sistec) if sistec is not None else "",
            "{CIDADE}": str(cidade) if cidade is not None else "",
            "{ESTADO}": str(estado) if estado is not None else ""
        }

        # Substituir todas as TAGs no documento
        substituir_todas_as_tags(doc, tags)

        # Verificar se houve algum problema no preenchimento (caso alguma TAG ainda esteja no documento)
        tags_nao_preenchidas = [tag for tag in tags if any(tag in paragraph.text for paragraph in doc.paragraphs)]

        # Verificar também dentro das tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        tags_nao_preenchidas += [tag for tag in tags if tag in paragraph.text]

        if tags_nao_preenchidas:
            erros_preenchimento.append({
                'aluno': nome_aluno,
                'tags': tags_nao_preenchidas
            })
            print(f"Erro: Algumas TAGs não foram preenchidas para {nome_aluno}: {tags_nao_preenchidas}")

        # Salvar o certificado preenchido
        try:
            doc.save(os.path.join(saida_certificados, f"{nome_aluno}_certificado.docx"))
            print(f"Certificado para {nome_aluno} gerado com sucesso!")
        except Exception as e:
            print(f"Erro ao gerar certificado para {nome_aluno}: {e}")

    # Relatório final de erros
    if erros_preenchimento:
        print("\nRelatório de TAGs não preenchidas:")
        for erro in erros_preenchimento:
            print(f"Aluno: {erro['aluno']} - TAGs não preenchidas: {erro['tags']}")
    else:
        print("Todos os certificados foram gerados com sucesso e todas as TAGs foram preenchidas!")


if __name__ == "__main__":
    planilha_excel = "alunos.xlsx"
    modelo_certificado = "modelo_certificado.docx"
    saida_certificados = "certificados"

    preencher_certificados(planilha_excel, modelo_certificado, saida_certificados)
