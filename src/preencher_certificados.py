from docx import Document
from openpyxl import load_workbook
import os
from datetime import datetime
import locale


def preencher_tags(paragraphs, tags_dict):
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for tag, value in tags_dict.items():
                if tag in run.text:
                    run.text = run.text.replace(tag, str(value))  # Garantir que o valor seja string


def preencher_tags_tabela(doc, tags_dict):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                preencher_tags(cell.paragraphs, tags_dict)


def converter_data(nascimento):
    """Converte diferentes formatos de data para o formato desejado."""
    meses = {
        1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
        5: "maio", 6: "junho", 7: "julho", 8: "agosto",
        9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
    }

    if isinstance(nascimento, datetime):
        nascimento_str = nascimento.strftime(f"%d de {meses[nascimento.month]} de %Y")
        return nascimento_str
    else:
        try:
            # Tentar converter string para datetime
            nascimento_dt = datetime.strptime(str(nascimento), "%d/%m/%Y")
            nascimento_str = nascimento_dt.strftime(f"%d de {meses[nascimento_dt.month]} de %Y")
            return nascimento_str
        except ValueError:
            # Se falhar, retornar a string original
            return str(nascimento)


def preencher_certificados(planilha_excel, modelo_certificado, saida_certificados):
    try:
        workbook = load_workbook(planilha_excel)
        planilha = workbook.active

        if not os.path.exists(saida_certificados):
            os.makedirs(saida_certificados)

        try:
            locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
        except locale.Error:
            print("Locale pt_BR.UTF-8 não disponível. Usando configuração padrão.")

        for row in range(2, planilha.max_row + 1):
            nome_aluno = planilha.cell(row=row, column=1).value or ""
            nascimento = planilha.cell(row=row, column=2).value
            cidade = planilha.cell(row=row, column=3).value or ""
            estado = planilha.cell(row=row, column=4).value or ""
            cpf = planilha.cell(row=row, column=5).value or ""
            n1 = planilha.cell(row=row, column=6).value
            n2 = planilha.cell(row=row, column=7).value
            n3 = planilha.cell(row=row, column=8).value
            n4 = planilha.cell(row=row, column=9).value
            n5 = planilha.cell(row=row, column=10).value
            n6 = planilha.cell(row=row, column=11).value
            n7 = planilha.cell(row=row, column=12).value
            n8 = planilha.cell(row=row, column=13).value
            n9 = planilha.cell(row=row, column=14).value
            n10 = planilha.cell(row=row, column=15).value
            reg = planilha.cell(row=row, column=16).value or ""
            sistec = planilha.cell(row=row, column=17).value or ""

            nascimento_str = converter_data(nascimento)
            print(f"Nascimento: {nascimento} -> Formatado: {nascimento_str}")  # Debug

            doc = Document(modelo_certificado)

            # Converter todos os valores para strings ao criar o dicionário de tags
            tags_dict = {
                "{NOME_ALUNO}": str(nome_aluno),
                "{NASCIMENTO}": str(nascimento_str),
                "{CIDADE}": str(cidade),
                "{ESTADO}": str(estado),
                "{CPF}": str(cpf),
                "{N1}": format(n1, ".1f") if n1 is not None else "",
                "{N2}": format(n2, ".1f") if n2 is not None else "",
                "{N3}": format(n3, ".1f") if n3 is not None else "",
                "{N4}": format(n4, ".1f") if n4 is not None else "",
                "{N5}": format(n5, ".1f") if n5 is not None else "",
                "{N6}": format(n6, ".1f") if n6 is not None else "",
                "{N7}": format(n7, ".1f") if n7 is not None else "",
                "{N8}": format(n8, ".1f") if n8 is not None else "",
                "{N9}": format(n9, ".1f") if n9 is not None else "",
                "{N10}": format(n10, ".1f") if n10 is not None else "",
                "{REG}": str(reg),
                "{SISTEC}": str(sistec)
            }

            preencher_tags(doc.paragraphs, tags_dict)
            preencher_tags_tabela(doc, tags_dict)

            nome_arquivo = f"{nome_aluno}_certificado.docx"
            caminho_arquivo = os.path.join(saida_certificados, nome_arquivo)
            doc.save(caminho_arquivo)
            print(f"Certificado preenchido para {nome_aluno} salvo em: {caminho_arquivo}")

    except Exception as e:
        print(f"Erro ao preencher certificados: {str(e)}")


if __name__ == "__main__":
    planilha_excel = "alunos.xlsx"
    modelo_certificado = "modelo_certificado.docx"
    saida_certificados = "certificados"

    preencher_certificados(planilha_excel, modelo_certificado, saida_certificados)
