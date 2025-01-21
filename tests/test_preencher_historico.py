import os
from docx import Document
from openpyxl import load_workbook


def carregar_dados(planilha):
    try:
        wb = load_workbook(planilha, data_only=True)
        ws = wb.active
        data = [row for row in ws.iter_rows(values_only=True)]
        print(f"Dados da planilha '{planilha}' carregados com sucesso.")
        return data
    except Exception as e:
        print(f"Erro ao carregar dados da planilha '{planilha}': {str(e)}")
        return None


def substituir_tags(doc_modelo, tags, dados, planilha_atual):
    def processar_run(run, tags, dados):
        for tag, value in zip(tags, dados):
            if tag in run.text:
                if value == "NDA" and planilha_atual in ["aluno_notas.xlsx", "aluno_conceito.xlsx"]:
                    value = " "
                elif tag.startswith("{N") and isinstance(value, (int, float)):
                    value = "{:.1f}".format(value)
                run.text = run.text.replace(tag, str(value))

    for paragraph in doc_modelo.paragraphs:
        for run in paragraph.runs:
            processar_run(run, tags, dados)

    for table in doc_modelo.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        processar_run(run, tags, dados)


def preencher_notas(planilhas, modelo_notas, saida_notas):
    dados_alunos, dados_notas, dados_conceito = map(
        carregar_dados, planilhas[:3])

    if any(dados is None for dados in
           [dados_alunos, dados_notas, dados_conceito]):
        print("Erro ao carregar dados das planilhas.")
        return

    os.makedirs(saida_notas, exist_ok=True)

    # Pega o índice da coluna REG para fazer a substituição correta
    header_alunos = dados_alunos[0]
    try:
        indice_reg = header_alunos.index("REG")
    except ValueError:
        print(
            "A coluna 'REG' não foi encontrada na planilha 'aluno_dados.xlsx'. Verifique se o cabeçalho está correto.")
        indice_reg = None  # Define um valor padrão para continuar sem a coluna REG

    for i, aluno_data in enumerate(dados_alunos[1:], start=1):
        # Atribuição de valores individuais
        nome_aluno, nascimento, nome_pai, nome_mae, cidade, uf = aluno_data[:6]
        reg = aluno_data[indice_reg] if indice_reg is not None else ""  # Valor da coluna REG ou string vazia

        # Verificações de valores vazios
        campos_obrigatorios = {
            "Nome do Aluno": nome_aluno,
            "Data de Nascimento": nascimento,
            "Nome do Pai": nome_pai,
            "Nome da Mãe": nome_mae,
            "Cidade": cidade,
            "UF": uf,
            "REG": reg
        }

        for campo, valor in campos_obrigatorios.items():
            if not valor:
                print(f"Aviso: o campo '{campo}' está vazio para o aluno '{nome_aluno}' na linha {i + 1} da planilha.")

        # Formata a data de nascimento caso não esteja vazia
        nascimento_formatado = nascimento.strftime("%d/%m/%Y") if nascimento else ""

        caminho_arquivo = os.path.join(saida_notas, f"{nome_aluno}_notas.docx")
        doc_modelo = Document(modelo_notas)

        substituir_tags(doc_modelo,
                        ["{NOME_ALUNO}", "{NASCIMENTO}", "{NOME_PAI}", "{NOME_MAE}", "{CIDADE}", "{UF}", "{REG}"],
                        [nome_aluno, nascimento_formatado, nome_pai, nome_mae, cidade, uf, reg],
                        planilhas[0])

        tags = [f"{{N{i}}}" for i in range(1, 21)] + [f"{{CON{i}}}" for i in range(1, 21)]

        dados_notas_aluno = dados_notas[i] if i < len(dados_notas) else [''] * 20
        dados_conceito_aluno = dados_conceito[i] if i < len(dados_conceito) else [''] * 20

        dados = dados_notas_aluno + dados_conceito_aluno

        substituir_tags(doc_modelo, tags, dados, planilhas[0])

        doc_modelo.save(caminho_arquivo)
        print(f"Arquivo {caminho_arquivo} gerado com sucesso!")

    print("Processo concluído!")


planilhas = [
    "aluno_dados.xlsx", "aluno_notas.xlsx", "aluno_conceito.xlsx"
]
saida_notas = "notas_alunos"

preencher_notas(planilhas, "modelo_historico.docx", saida_notas)
